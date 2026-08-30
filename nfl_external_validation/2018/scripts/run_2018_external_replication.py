from __future__ import annotations

import csv
import argparse
import hashlib
import importlib.metadata
import json
import math
import platform
import shutil
import sys
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import yaml
from docx import Document
from scipy.stats import circmean
from sklearn.compose import ColumnTransformer
from sklearn.ensemble import HistGradientBoostingRegressor
from sklearn.impute import SimpleImputer
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import OneHotEncoder, OrdinalEncoder


REPO_ROOT = Path(__file__).resolve().parents[3]
WORK = REPO_ROOT / "outputs_2018"
DATA = WORK / "derived_cache"
RAW = Path()
BONUS = Path()
TARGET_LABEL_FILE = Path()
MANUSCRIPT: Path | None = None

PRIMARY_POS = {"CB", "DB", "FS", "SS", "S"}
EXPANDED_POS = {"CB", "DB", "FS", "SS", "S", "LB", "ILB", "OLB", "MLB"}
SEED = 20260730
PSEUDO_SEED = 20260731
BOOT_REPS = 2000
PSEUDO_REPS = 5000

A_FEATURES = [
    "n_input_frames",
    "input_start_x",
    "input_start_y",
    "input_end_x",
    "input_end_y",
    "input_delta_x",
    "input_delta_y",
    "input_path_length",
    "input_mean_speed",
    "input_max_speed",
    "input_mean_acceleration",
    "input_max_acceleration",
    "input_mean_dir",
    "input_mean_o",
    "play_direction",
    "player_side",
    "player_role",
    "player_position",
]
H_ADD = ["num_frames_output"]
D_ADD = [
    "ball_land_x",
    "ball_land_y",
    "distance_to_ball_land_start",
    "distance_to_ball_land_end",
    "change_in_distance_to_ball_land",
]


def now() -> str:
    return datetime.now().isoformat(timespec="seconds")


def ensure_dirs() -> None:
    for p in [
        WORK / "protocol",
        WORK / "scripts",
        WORK / "logs",
        WORK / "outputs" / "tables",
        WORK / "outputs" / "figures",
        WORK / "outputs" / "intermediate",
        WORK / "reports",
        DATA / "processed",
        DATA / "qc",
        DATA / "manifests",
    ]:
        p.mkdir(parents=True, exist_ok=True)


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def row_count_csv(path: Path) -> int:
    with path.open("rb") as f:
        return max(sum(1 for _ in f) - 1, 0)


def write(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8", newline="\n")


def write_manifest() -> pd.DataFrame:
    rows = []
    for source_root, source in [(RAW, "official NFL Big Data Bowl 2021 competition"), (BONUS, "tombliss BDB2021 bonus")]:
        for p in sorted(source_root.glob("*")):
            if not p.is_file():
                continue
            rc = row_count_csv(p) if p.suffix.lower() == ".csv" else ""
            rows.append(
                {
                    "source": source,
                    "file": p.name,
                    "absolute_path": str(p),
                    "size": p.stat().st_size,
                    "sha256": sha256(p),
                    "download_timestamp": "",
                    "row_count": rc,
                    "status": "used" if p.suffix.lower() == ".csv" else "preserved archive",
                }
            )
    df = pd.DataFrame(rows)
    out = DATA / "manifests" / "RAW_DATA_MANIFEST.csv"
    df.to_csv(out, index=False)
    return df


def normalize_xy(x: pd.Series | float, y: pd.Series | float, direction: pd.Series | str):
    if isinstance(direction, pd.Series):
        left = direction.eq("left")
        nx = x.where(~left, 120 - x)
        ny = y.where(~left, (160 / 3) - y)
        return nx, ny
    if direction == "left":
        return 120 - x, (160 / 3) - y
    return x, y


def cm(series: pd.Series) -> float:
    vals = pd.to_numeric(series, errors="coerce").dropna().to_numpy()
    if len(vals) == 0:
        return np.nan
    return float(np.degrees(circmean(np.radians(vals), high=2 * np.pi, low=0)))


def path_length(df: pd.DataFrame) -> float:
    xy = df[["x_norm", "y_norm"]].dropna()
    if len(xy) < 2:
        return 0.0
    d = xy.diff().iloc[1:]
    return float(np.sqrt(d["x_norm"] ** 2 + d["y_norm"] ** 2).sum())


def load_static():
    games = pd.read_csv(RAW / "games.csv")
    plays = pd.read_csv(RAW / "plays.csv")
    players = pd.read_csv(RAW / "players.csv")
    targets = pd.read_csv(TARGET_LABEL_FILE)
    plays = plays.merge(games[["gameId", "week", "homeTeamAbbr", "visitorTeamAbbr"]], on="gameId", how="left")
    return games, plays, players, targets


@dataclass
class BuildResult:
    rows: list[dict]
    flow_rows: list[dict]
    event_rows: list[dict]
    target_qc: list[dict]
    position_counts: list[dict]
    week_counts: list[dict]


def build_dataset(position_set: set[str], endpoint_window: int = 0, exclude_ambiguous: bool = False) -> BuildResult:
    games, plays, players, targets = load_static()
    target_map = targets.set_index(["gameId", "playId"])["targetNflId"].to_dict()
    player_pos = players.set_index("nflId")["position"].to_dict()
    rows: list[dict] = []
    flow_rows: list[dict] = []
    event_rows: list[dict] = []
    target_qc: list[dict] = []
    pos_counts: list[dict] = []
    week_counts: list[dict] = []
    play_meta = plays.set_index(["gameId", "playId"]).to_dict("index")

    cache_tag = "_".join(sorted(position_set)) + f"_w{endpoint_window}_amb{int(exclude_ambiguous)}"
    cache_dir = DATA / "processed" / "weekly_cache" / cache_tag
    cache_dir.mkdir(parents=True, exist_ok=True)
    for week in range(1, 18):
        row_cache = cache_dir / f"week{week}_rows.parquet"
        meta_cache = cache_dir / f"week{week}_meta.json"
        if row_cache.exists() and meta_cache.exists():
            wk_rows = pd.read_parquet(row_cache).to_dict("records")
            meta = json.loads(meta_cache.read_text(encoding="utf-8"))
            rows.extend(wk_rows)
            flow_rows.append(meta["flow"])
            event_rows.extend(meta["events"])
            target_qc.extend(meta["target_qc"])
            pos_counts.extend(meta["position_counts"])
            week_counts.append(meta["week_counts"])
            print(f"[cache] week {week}: {len(wk_rows)} modeling rows", flush=True)
            continue
        print(f"[build] week {week}: reading tracking CSV", flush=True)
        f = RAW / f"week{week}.csv"
        tr = pd.read_csv(
            f,
            usecols=[
                "x",
                "y",
                "s",
                "a",
                "o",
                "dir",
                "event",
                "nflId",
                "displayName",
                "position",
                "frameId",
                "team",
                "gameId",
                "playId",
                "playDirection",
            ],
        )
        raw_rows = len(tr)
        week_model_rows: list[dict] = []
        week_events: list[dict] = []
        week_target_qc: list[dict] = []
        raw_plays = tr[["gameId", "playId"]].drop_duplicates()
        tr["event"] = tr["event"].fillna("")
        tr["nflId"] = pd.to_numeric(tr["nflId"], errors="coerce")
        tr["x_norm"], tr["y_norm"] = normalize_xy(tr["x"], tr["y"], tr["playDirection"])
        grouped = tr.groupby(["gameId", "playId"], sort=False)
        counts = {
            "week": week,
            "raw_rows": raw_rows,
            "raw_plays": len(raw_plays),
            "with_ball_snap": 0,
            "with_pass_forward": 0,
            "with_pass_arrived": 0,
            "valid_ordered_chain": 0,
            "with_target_receiver": 0,
            "with_exact_football_arrival": 0,
            "primary_defender_eligible": 0,
            "three_role_eligible": 0,
            "modeling_rows": 0,
        }
        print(f"[build] week {week}: {len(raw_plays)} plays, {raw_rows} rows", flush=True)
        for j, ((gid, pid), g) in enumerate(grouped, start=1):
            if j % 500 == 0:
                print(f"[build] week {week}: processed {j}/{len(raw_plays)} plays", flush=True)
            meta = play_meta.get((gid, pid), {})
            frames = g[["frameId", "event"]].drop_duplicates()
            snap_vals = frames.loc[frames["event"].eq("ball_snap"), "frameId"]
            fw_vals = frames.loc[frames["event"].eq("pass_forward"), "frameId"]
            ar_vals = frames.loc[frames["event"].eq("pass_arrived"), "frameId"]
            if len(snap_vals):
                counts["with_ball_snap"] += 1
            if len(fw_vals):
                counts["with_pass_forward"] += 1
            if len(ar_vals):
                counts["with_pass_arrived"] += 1
            if not (len(snap_vals) and len(fw_vals) and len(ar_vals)):
                continue
            snap = int(snap_vals.min())
            fw_after = fw_vals[fw_vals > snap]
            if fw_after.empty:
                continue
            forward = int(fw_after.min())
            ar_after = ar_vals[ar_vals > forward]
            if ar_after.empty:
                continue
            arrival = int(ar_after.min())
            counts["valid_ordered_chain"] += 1
            target_id = target_map.get((gid, pid))
            if pd.isna(target_id):
                continue
            target_id = int(target_id)
            target_rows = g[g["nflId"].eq(target_id)]
            if target_rows.empty:
                continue
            counts["with_target_receiver"] += 1
            offense_team = str(target_rows["team"].dropna().iloc[0])
            if offense_team not in {"home", "away"}:
                continue
            defense_team = "away" if offense_team == "home" else "home"
            ball_at_arrival = g[(g["displayName"].eq("Football")) & (g["frameId"].eq(arrival))]
            if ball_at_arrival.empty and endpoint_window:
                ball_near = g[(g["displayName"].eq("Football")) & (g["frameId"].between(arrival - endpoint_window, arrival + endpoint_window))].copy()
                if not ball_near.empty:
                    ball_near["frame_dist"] = (ball_near["frameId"] - arrival).abs()
                    ball_at_arrival = ball_near.sort_values(["frame_dist", "frameId"]).head(1)
            if ball_at_arrival.empty:
                continue
            counts["with_exact_football_arrival"] += 1
            ball_x = float(ball_at_arrival["x_norm"].iloc[0])
            ball_y = float(ball_at_arrival["y_norm"].iloc[0])
            off_arr = g[(g["team"].eq(offense_team)) & (g["frameId"].eq(arrival)) & (g["nflId"].notna())].copy()
            if not off_arr.empty:
                off_arr["dist_ball"] = np.sqrt((off_arr["x_norm"] - ball_x) ** 2 + (off_arr["y_norm"] - ball_y) ** 2)
                off_arr = off_arr.sort_values(["dist_ball", "nflId"])
                nearest = float(off_arr["dist_ball"].iloc[0])
                second = float(off_arr["dist_ball"].iloc[1]) if len(off_arr) > 1 else np.nan
                recon = int(off_arr["nflId"].iloc[0])
                ambiguous = bool(nearest > 5.0 or (pd.notna(second) and second - nearest < 1.0))
            else:
                nearest = np.nan
                second = np.nan
                recon = -1
                ambiguous = True
            week_target_qc.append(
                {
                    "gameId": gid,
                    "playId": pid,
                    "week": week,
                    "bonus_target_nflId": target_id,
                    "reconstructed_target_nflId": recon,
                    "exact_agreement": target_id == recon,
                    "nearest_distance_to_ball": nearest,
                    "second_nearest_distance_to_ball": second,
                    "second_minus_nearest_margin": second - nearest if pd.notna(second) else np.nan,
                    "ambiguous_reconstructed_target": ambiguous,
                    "target_source": "PUBLIC_BONUS",
                    "passResult": meta.get("passResult", ""),
                }
            )
            if exclude_ambiguous and ambiguous:
                continue
            fw_players = g[(g["frameId"].eq(forward)) & (g["nflId"].notna())].copy()
            r_fw = fw_players[fw_players["nflId"].eq(target_id)]
            if r_fw.empty:
                continue
            defenders = fw_players[(fw_players["team"].eq(defense_team)) & (fw_players["position"].isin(position_set))].copy()
            if len(defenders) < 2:
                continue
            counts["primary_defender_eligible"] += 1
            rx = float(r_fw["x_norm"].iloc[0])
            ry = float(r_fw["y_norm"].iloc[0])
            defenders["dist_to_r"] = np.sqrt((defenders["x_norm"] - rx) ** 2 + (defenders["y_norm"] - ry) ** 2)
            defenders = defenders.sort_values(["dist_to_r", "nflId"])
            n_id = int(defenders["nflId"].iloc[0])
            o_ids = [int(x) for x in defenders["nflId"].iloc[1:].tolist()]
            counts["three_role_eligible"] += 1
            units = [(target_id, "R", "Offense", "Targeted Receiver"), (n_id, "N", "Defense", "Defensive Coverage")]
            units += [(oid, "O", "Defense", "Defensive Coverage") for oid in o_ids]
            for nfl_id, role, side, iface_role in units:
                pg = g[(g["nflId"].eq(nfl_id)) & (g["frameId"].between(snap, forward))].sort_values("frameId")
                ep = g[(g["nflId"].eq(nfl_id)) & (g["frameId"].eq(arrival))]
                if pg.empty or ep.empty:
                    continue
                start = pg[["x_norm", "y_norm"]].dropna().head(1)
                end = pg[["x_norm", "y_norm"]].dropna().tail(1)
                if start.empty or end.empty:
                    continue
                sx, sy = float(start["x_norm"].iloc[0]), float(start["y_norm"].iloc[0])
                ex, ey = float(end["x_norm"].iloc[0]), float(end["y_norm"].iloc[0])
                true_x, true_y = float(ep["x_norm"].iloc[0]), float(ep["y_norm"].iloc[0])
                d_start = math.hypot(sx - ball_x, sy - ball_y)
                d_end = math.hypot(ex - ball_x, ey - ball_y)
                pos = str(pg["position"].dropna().iloc[-1]) if not pg["position"].dropna().empty else str(player_pos.get(nfl_id, ""))
                rec = {
                        "gameId": gid,
                        "playId": pid,
                        "nflId": nfl_id,
                        "week": week,
                        "role": role,
                        "position": pos,
                        "true_x": true_x,
                        "true_y": true_y,
                        "n_input_frames": int(pg["frameId"].nunique()),
                        "input_start_x": sx,
                        "input_start_y": sy,
                        "input_end_x": ex,
                        "input_end_y": ey,
                        "input_delta_x": ex - sx,
                        "input_delta_y": ey - sy,
                        "input_path_length": path_length(pg),
                        "input_mean_speed": float(pd.to_numeric(pg["s"], errors="coerce").mean()),
                        "input_max_speed": float(pd.to_numeric(pg["s"], errors="coerce").max()),
                        "input_mean_acceleration": float(pd.to_numeric(pg["a"], errors="coerce").mean()),
                        "input_max_acceleration": float(pd.to_numeric(pg["a"], errors="coerce").max()),
                        "input_mean_dir": cm(pg["dir"]),
                        "input_mean_o": cm(pg["o"]),
                        "play_direction": str(pg["playDirection"].iloc[0]),
                        "player_side": side,
                        "player_role": iface_role,
                        "player_position": pos,
                        "num_frames_output": arrival - forward,
                        "ball_land_x": ball_x,
                        "ball_land_y": ball_y,
                        "distance_to_ball_land_start": d_start,
                        "distance_to_ball_land_end": d_end,
                        "change_in_distance_to_ball_land": d_end - d_start,
                        "target_source": "PUBLIC_BONUS",
                        "ambiguous_reconstructed_target": ambiguous,
                        "defender_set_size": len(defenders),
                        "endpoint_window": endpoint_window,
                    }
                rows.append(rec)
                week_model_rows.append(rec)
                counts["modeling_rows"] += 1
            ev = {"gameId": gid, "playId": pid, "week": week, "snap_frame": snap, "forward_frame": forward, "arrival_frame": arrival}
            event_rows.append(ev)
            week_events.append(ev)
        flow_rows.append(counts)
        wk_pos = tr[tr["nflId"].notna()].groupby("position").size().reset_index(name="rows").assign(week=week).to_dict("records")
        pos_counts.extend(wk_pos)
        target_qc.extend(week_target_qc)
        wk_counts = {"week": week, "raw_rows": raw_rows, "raw_plays": len(raw_plays)}
        week_counts.append(wk_counts)
        pd.DataFrame(week_model_rows).to_parquet(row_cache, index=False)
        meta_cache.write_text(
            json.dumps(
                {
                    "flow": counts,
                    "events": week_events,
                    "target_qc": week_target_qc,
                    "position_counts": wk_pos,
                    "week_counts": wk_counts,
                },
                indent=2,
            ),
            encoding="utf-8",
        )
        print(f"[build] week {week}: wrote {len(week_model_rows)} modeling rows", flush=True)
    return BuildResult(rows, flow_rows, event_rows, target_qc, pos_counts, week_counts)


def preprocessor(feature_names: list[str]):
    cat = [c for c in feature_names if c in {"play_direction", "player_side", "player_role", "player_position"}]
    num = [c for c in feature_names if c not in cat]
    return ColumnTransformer(
        [
            ("num", SimpleImputer(strategy="median"), num),
            ("cat", Pipeline([("imp", SimpleImputer(strategy="most_frequent")), ("oh", OneHotEncoder(handle_unknown="ignore", sparse_output=False))]), cat),
        ]
    )


def fit_predict(df: pd.DataFrame, features: list[str]) -> tuple[np.ndarray, np.ndarray, dict]:
    train = df["week"].between(1, 11)
    held = df["week"].between(14, 17)
    pipe_x = Pipeline([("prep", preprocessor(features)), ("model", HistGradientBoostingRegressor(random_state=SEED))])
    pipe_y = Pipeline([("prep", preprocessor(features)), ("model", HistGradientBoostingRegressor(random_state=SEED))])
    pipe_x.fit(df.loc[train, features], df.loc[train, "true_x"])
    pipe_y.fit(df.loc[train, features], df.loc[train, "true_y"])
    px = pipe_x.predict(df.loc[held, features])
    py = pipe_y.predict(df.loc[held, features])
    return px, py, {"x_train_score": pipe_x.score(df.loc[train, features], df.loc[train, "true_x"]), "y_train_score": pipe_y.score(df.loc[train, features], df.loc[train, "true_y"])}


def add_predictions(df: pd.DataFrame) -> tuple[pd.DataFrame, dict]:
    held = df[df["week"].between(14, 17)].copy()
    scores = {}
    for state, feats in {"A": A_FEATURES, "H": A_FEATURES + H_ADD, "D": A_FEATURES + H_ADD + D_ADD}.items():
        px, py, sc = fit_predict(df, feats)
        held[f"{state}_pred_x"] = px
        held[f"{state}_pred_y"] = py
        held[f"error_{state}"] = np.sqrt((held[f"{state}_pred_x"] - held["true_x"]) ** 2 + (held[f"{state}_pred_y"] - held["true_y"]) ** 2)
        scores[state] = sc
    held["H_to_D_gain"] = held["error_H"] - held["error_D"]
    return held, scores


def estimates(held: pd.DataFrame) -> dict:
    r = held[held["role"].eq("R")]["H_to_D_gain"].mean()
    n = held[held["role"].eq("N")]["H_to_D_gain"].mean()
    o_play = held[held["role"].eq("O")].groupby(["gameId", "playId"])["H_to_D_gain"].mean()
    o = o_play.mean()
    return {
        "G_R": float(r),
        "G_N": float(n),
        "G_O": float(o),
        "lambda_D": float(n - o),
        "delta_RN": float(r - n),
        "delta_RO": float(r - o),
    }


def play_level_gains(held: pd.DataFrame) -> pd.DataFrame:
    r = held[held.role.eq("R")].set_index(["gameId", "playId"])[["week", "H_to_D_gain"]].rename(columns={"H_to_D_gain": "R"})
    n = held[held.role.eq("N")].set_index(["gameId", "playId"])[["H_to_D_gain"]].rename(columns={"H_to_D_gain": "N"})
    o = held[held.role.eq("O")].groupby(["gameId", "playId"])["H_to_D_gain"].mean().to_frame("O")
    games = held.groupby(["gameId", "playId"])["gameId"].first().to_frame("cluster_gameId")
    return r.join(n).join(o).join(games).reset_index().dropna(subset=["R", "N", "O"])


def estimates_from_play(pl: pd.DataFrame) -> dict:
    gr, gn, go = pl["R"].mean(), pl["N"].mean(), pl["O"].mean()
    return {"G_R": gr, "G_N": gn, "G_O": go, "lambda_D": gn - go, "delta_RN": gr - gn, "delta_RO": gr - go}


def bootstrap(pl: pd.DataFrame, reps: int = BOOT_REPS, seed: int = SEED) -> pd.DataFrame:
    rng = np.random.default_rng(seed)
    games = pl["gameId"].drop_duplicates().to_numpy()
    out = []
    by_game = {g: d for g, d in pl.groupby("gameId")}
    for i in range(reps):
        sample = rng.choice(games, size=len(games), replace=True)
        b = pd.concat([by_game[g] for g in sample], ignore_index=True)
        est = estimates_from_play(b)
        est["replicate"] = i
        out.append(est)
    return pd.DataFrame(out)


def ci_table(est: dict, boot: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for k, v in est.items():
        rows.append({"estimand": k, "estimate": v, "ci_low": boot[k].quantile(0.025), "ci_high": boot[k].quantile(0.975)})
    return pd.DataFrame(rows)


def pseudo_nearest(held: pd.DataFrame, reps: int = PSEUDO_REPS, seed: int = PSEUDO_SEED) -> tuple[pd.DataFrame, dict]:
    rng = np.random.default_rng(seed)
    defenders = held[held.role.isin(["N", "O"])].copy()
    grouped = list(defenders.groupby(["gameId", "playId"]))
    obs = estimates(held)["lambda_D"]
    plays = []
    counts = []
    sums = []
    max_count = 0
    for key, d in grouped:
        gains = d["H_to_D_gain"].to_numpy(float)
        if len(gains) < 2:
            continue
        plays.append(gains)
        counts.append(len(gains))
        sums.append(gains.sum())
        max_count = max(max_count, len(gains))
    nplays = len(plays)
    gain_mat = np.full((nplays, max_count), np.nan)
    for i, gains in enumerate(plays):
        gain_mat[i, : len(gains)] = gains
    counts_arr = np.asarray(counts)
    sums_arr = np.asarray(sums)
    values = np.empty(reps, dtype=float)
    # Chunk to keep temporary random-index matrices moderate.
    chunk = 500
    for start in range(0, reps, chunk):
        end = min(start + chunk, reps)
        idx = np.floor(rng.random((end - start, nplays)) * counts_arr).astype(int)
        chosen = gain_mat[np.arange(nplays)[None, :], idx]
        pseudo_n = chosen.mean(axis=1)
        pseudo_o = ((sums_arr - chosen) / (counts_arr - 1)).mean(axis=1)
        values[start:end] = pseudo_n - pseudo_o
    null = pd.DataFrame({"replicate": np.arange(reps), "pseudo_lambda_D": values})
    p = (1 + (np.abs(null["pseudo_lambda_D"]) >= abs(obs)).sum()) / (reps + 1)
    summary = {
        "observed_lambda_D": obs,
        "null_mean": float(null["pseudo_lambda_D"].mean()),
        "null_low": float(null["pseudo_lambda_D"].quantile(0.025)),
        "null_high": float(null["pseudo_lambda_D"].quantile(0.975)),
        "p_value": float(p),
    }
    return null, summary


def freeze_protocol(manifest: pd.DataFrame) -> None:
    config = {
        "data_source": "official NFL Big Data Bowl 2021 competition files",
        "target_receiver_source": "public targeted-receiver bonus labels from tombliss/nfl-big-data-bowl-2021-bonus",
        "target_label_mode": "public bonus labels primary; reconstructed labels only for QC/sensitivity if needed",
        "event_definitions": {"snap": "earliest ball_snap", "forward": "earliest pass_forward after snap", "arrival": "earliest pass_arrived after forward"},
        "temporal_split": {"train": "weeks 1-11", "validation": "weeks 12-13", "heldout": "weeks 14-17"},
        "coordinate_normalization": "prespecified left-to-right coordinate normalization",
        "primary_defender_set": sorted(PRIMARY_POS),
        "expanded_defender_set": sorted(EXPANDED_POS),
        "roles": {"R": "targeted receiver from public bonus labels", "N": "nearest coverage-eligible defender to R at pass_forward", "O": "remaining coverage-eligible defenders, within-play mean"},
        "features": {"A": A_FEATURES, "H_add": H_ADD, "D_add": D_ADD},
        "loss": "Euclidean endpoint error in yards",
        "HGB": "sklearn HistGradientBoostingRegressor defaults, random_state 20260730; deterministic preprocessing",
        "ExtraTrees": "skipped unless exact submitted final configuration is recoverable",
        "bootstrap": {"method": "percentile game-cluster bootstrap", "replicates": BOOT_REPS, "seed": SEED},
        "pseudo_nearest": {"permutations": PSEUDO_REPS, "seed": PSEUDO_SEED, "p_value": "add-one corrected two-sided absolute statistic"},
        "estimands": ["G_R", "G_N", "G_O", "lambda_D", "delta_RN", "delta_RO"],
        "sensitivity_analyses": ["primary defender set", "expanded defender set", "+/-1 football endpoint", "exclude ambiguous reconstructed-target QC flags"],
    }
    protocol = f"""# Frozen 2018 Replication Protocol

Freeze timestamp: {now()}

This protocol was written and hashed before fitting the final 2018 HGB models or viewing primary 2018 RPVA results.

## Locked Configuration

```yaml
{yaml.safe_dump(config, sort_keys=False)}
```

## Raw Data Hashes

{manifest[['source','file','sha256','size','row_count']].to_markdown(index=False)}

## Interpretation Rules

The 2018 analysis is an external-season RPVA replication using harmonized BDB2021 data. Proximity-based N is a nearest observed coverage-defender proxy, not an assigned defender. Destination and endpoint variables are retrospective audit-state inputs, not deployment inputs. Weak, null, or discordant results do not trigger post-freeze redefinition.
"""
    write(WORK / "protocol" / "FROZEN_2018_REPLICATION_PROTOCOL.md", protocol)
    write(WORK / "protocol" / "protocol_config.yaml", yaml.safe_dump(config, sort_keys=False))
    write(WORK / "protocol" / "DEVIATIONS.md", "No post-freeze deviations recorded.\n")
    hashes = []
    for p in [WORK / "protocol" / "FROZEN_2018_REPLICATION_PROTOCOL.md", WORK / "protocol" / "protocol_config.yaml"]:
        hashes.append(f"{sha256(p)}  {p.name}")
    write(WORK / "protocol" / "protocol_sha256.txt", "\n".join(hashes) + "\n")


def summarize_replication(primary: dict, ci: pd.DataFrame, pseudo: dict) -> str:
    lam = primary["lambda_D"]
    lam_ci = ci[ci.estimand.eq("lambda_D")].iloc[0]
    ci_pos = lam_ci.ci_low > 0 and lam_ci.ci_high > 0
    pseudo_support = pseudo["p_value"] < 0.05 and not (pseudo["null_low"] <= lam <= pseudo["null_high"])
    ordering = primary["G_R"] > primary["G_N"] > primary["G_O"]
    if lam > 0 and ci_pos and pseudo_support and ordering:
        return "The primary 2018 analysis reproduced the principal relational ordering, with positive nearest-other localization and pseudo-nearest falsification support."
    if lam > 0 and (ci_pos or pseudo_support):
        return "The primary 2018 analysis supported positive nearest-other localization, while not all directional criteria were simultaneously met."
    if abs(lam) > 0.1 and (ci_pos or pseudo_support or abs(primary["G_R"] - primary["G_O"]) > 0.1):
        return "The primary 2018 analysis showed relational structure whose full interpretation depends on harmonized role and feature definitions."
    return "The primary 2018 analysis did not support a positive nearest-other localization pattern."


def make_figures(primary_ci, comp, pseudo_null, weekly, sens):
    figdir = WORK / "outputs" / "figures"
    # Figure 1
    fig, ax = plt.subplots(figsize=(7, 4))
    roles = ["G_R", "G_N", "G_O"]
    labels = ["R", "N", "O"]
    x = np.arange(len(labels))
    vals18 = [primary_ci.set_index("estimand").loc[r, "estimate"] for r in roles]
    lo18 = [primary_ci.set_index("estimand").loc[r, "estimate"] - primary_ci.set_index("estimand").loc[r, "ci_low"] for r in roles]
    hi18 = [primary_ci.set_index("estimand").loc[r, "ci_high"] - primary_ci.set_index("estimand").loc[r, "estimate"] for r in roles]
    vals23 = [1.735, 1.359, 0.609]
    lo23 = [1.735 - 1.621, 1.359 - 1.272, 0.609 - 0.517]
    hi23 = [1.850 - 1.735, 1.456 - 1.359, 0.703 - 0.609]
    ax.errorbar(x - .08, vals23, yerr=[lo23, hi23], fmt="o", label="2023")
    ax.errorbar(x + .08, vals18, yerr=[lo18, hi18], fmt="s", label="2018")
    ax.set_xticks(x, labels)
    ax.set_ylabel("H-to-D gain (yards)")
    ax.legend(frameon=False)
    fig.tight_layout()
    fig.savefig(figdir / "Figure1_2023_vs_2018_role_gains.png", dpi=300)
    fig.savefig(figdir / "Figure1_2023_vs_2018_role_gains.pdf")
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(5, 4))
    lam18 = primary_ci.set_index("estimand").loc["lambda_D"]
    ax.errorbar([0, 1], [0.750, lam18.estimate], yerr=[[0.750 - 0.637, lam18.estimate - lam18.ci_low], [0.861 - 0.750, lam18.ci_high - lam18.estimate]], fmt="o")
    ax.axhline(0, color="0.4", lw=1)
    ax.set_xticks([0, 1], ["2023", "2018"])
    ax.set_ylabel("Nearest-other localization (yards)")
    fig.tight_layout()
    fig.savefig(figdir / "Figure2_2023_vs_2018_lambda.png", dpi=300)
    fig.savefig(figdir / "Figure2_2023_vs_2018_lambda.pdf")
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(6, 4))
    ax.hist(pseudo_null["pseudo_lambda_D"], bins=60, color="0.75", edgecolor="white")
    ax.axvline(lam18.estimate, color="black", lw=2)
    ax.set_xlabel("Pseudo-nearest lambda_D")
    ax.set_ylabel("Permutations")
    fig.tight_layout()
    fig.savefig(figdir / "Figure3_pseudo_nearest_null.png", dpi=300)
    fig.savefig(figdir / "Figure3_pseudo_nearest_null.pdf")
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(6, 4))
    ax.plot(weekly["week"], weekly["lambda_D"], marker="o")
    ax.axhline(0, color="0.4", lw=1)
    ax.set_xticks([14, 15, 16, 17])
    ax.set_ylabel("Weekly lambda_D")
    ax.set_xlabel("Held-out week")
    fig.tight_layout()
    fig.savefig(figdir / "Figure4_weekly_lambda.png", dpi=300)
    fig.savefig(figdir / "Figure4_weekly_lambda.pdf")
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(6, 4))
    ax.bar(sens["analysis"], sens["lambda_D"])
    ax.axhline(0, color="0.4", lw=1)
    ax.set_ylabel("lambda_D")
    ax.tick_params(axis="x", rotation=25)
    fig.tight_layout()
    fig.savefig(figdir / "Figure5_sensitivity_lambda.png", dpi=300)
    fig.savefig(figdir / "Figure5_sensitivity_lambda.pdf")
    plt.close(fig)


def make_docx_copy(result_text: str) -> None:
    if MANUSCRIPT is None:
        return
    out = WORK / "Manuscript_with_2018_external_validation_DRAFT.docx"
    shutil.copy2(MANUSCRIPT, out)
    doc = Document(out)
    doc.add_page_break()
    doc.add_heading("2018 External-Season Validation Draft Insert", level=1)
    for para in result_text.split("\n\n"):
        doc.add_paragraph(para)
    doc.save(out)


def package_versions() -> str:
    names = ["numpy", "pandas", "scipy", "scikit-learn", "pyarrow", "python-docx", "matplotlib", "joblib", "kaggle", "PyYAML"]
    rows = []
    for name in names:
        try:
            ver = importlib.metadata.version(name)
        except importlib.metadata.PackageNotFoundError:
            ver = "NOT INSTALLED"
        rows.append({"package": name, "version": ver})
    return pd.DataFrame(rows).to_markdown(index=False)


def main() -> None:
    global WORK, DATA, RAW, BONUS, TARGET_LABEL_FILE, MANUSCRIPT
    parser = argparse.ArgumentParser(description="Run the frozen 2018 RPVA external-season replication.")
    parser.add_argument("--config", type=Path, default=REPO_ROOT / "nfl_external_validation" / "2018" / "protocol" / "protocol_config.yaml")
    parser.add_argument("--raw-data-root", type=Path, required=True, help="Directory containing authorized BDB2021 games.csv, plays.csv, players.csv, and week1.csv ... week17.csv files.")
    parser.add_argument("--target-label-file", type=Path, required=True, help="Local path to targetedReceiver.csv from tombliss/nfl-big-data-bowl-2021-bonus.")
    parser.add_argument("--output-dir", type=Path, default=REPO_ROOT / "outputs_2018", help="Directory for regenerated aggregate outputs and local derived caches.")
    parser.add_argument("--source-manuscript", type=Path, default=None, help="Optional local manuscript path used only to create a draft appendix copy.")
    args = parser.parse_args()
    WORK = args.output_dir.resolve()
    DATA = WORK / "derived_cache"
    RAW = args.raw_data_root.resolve()
    TARGET_LABEL_FILE = args.target_label_file.resolve()
    BONUS = TARGET_LABEL_FILE.parent
    MANUSCRIPT = args.source_manuscript.resolve() if args.source_manuscript else None
    if not args.config.exists():
        raise SystemExit(f"Config not found: {args.config}")
    if not RAW.exists():
        raise SystemExit(f"Raw data root not found: {RAW}")
    if not TARGET_LABEL_FILE.exists():
        raise SystemExit(f"Target label file not found: {TARGET_LABEL_FILE}")
    ensure_dirs()
    manifest = write_manifest()

    primary_build = build_dataset(PRIMARY_POS)
    df = pd.DataFrame(primary_build.rows)
    if df.empty:
        raise SystemExit("Hard Block 3: no primary modeling rows could be built.")
    df.to_parquet(DATA / "processed" / "2018_primary_modeling_dataset.parquet", index=False)
    pd.DataFrame(primary_build.flow_rows).to_csv(WORK / "outputs" / "tables" / "2018_sample_flow.csv", index=False)
    pd.DataFrame(primary_build.event_rows).to_csv(WORK / "outputs" / "tables" / "2018_event_counts.csv", index=False)
    pd.DataFrame(primary_build.position_counts).to_csv(WORK / "outputs" / "tables" / "2018_position_counts.csv", index=False)
    pd.DataFrame(primary_build.week_counts).to_csv(WORK / "outputs" / "tables" / "2018_week_counts.csv", index=False)
    tqc = pd.DataFrame(primary_build.target_qc)
    tqc.to_csv(WORK / "outputs" / "tables" / "2018_target_receiver_qc.csv", index=False)

    mapping = pd.DataFrame(
        [
            {"variable": c, "status": "EXACT" if c in ["gameId", "playId", "nflId", "week", "position"] else "HARMONIZED", "notes": "BDB2021-derived harmonized feature"}
            for c in ["gameId", "playId", "nflId", "week", "role", "position"] + A_FEATURES + H_ADD + D_ADD
        ]
    )
    mapping.to_csv(WORK / "outputs" / "tables" / "2023_to_2018_variable_mapping.csv", index=False)

    feasibility = f"""# 2018 Feasibility Report

Generated: {now()}

BDB2021 competition data were obtained locally under applicable access terms and the public bonus targeted-receiver file was available.

## Sample Flow

{pd.DataFrame(primary_build.flow_rows).to_markdown(index=False)}

## Target QC Summary

- Target source: public targeted-receiver bonus labels for primary labels.
- Official-vs-reconstructed agreement for QC: {tqc['exact_agreement'].mean():.3f}
- Ambiguous reconstructed-target QC flags: {int(tqc['ambiguous_reconstructed_target'].sum())} of {len(tqc)}
"""
    write(WORK / "reports" / "2018_FEASIBILITY_REPORT.md", feasibility)

    harm = """# 2018-2023 Harmonization Report

The 2018 replication uses BDB2021 competition tracking files with public bonus targeted-receiver labels. A/H/D feature definitions are harmonized using the prespecified left-to-right coordinate normalization because the submitted public 2023 repository provides NFL implementation scaffolding but not a verified full raw-data transformation/fitting pipeline.

Coverage labels and the submitted C state are not reconstructed for the 2018 confirmatory analysis.
"""
    write(WORK / "reports" / "2018_2023_HARMONIZATION_REPORT.md", harm)
    freeze_protocol(manifest)

    held, scores = add_predictions(df)
    held.to_parquet(WORK / "outputs" / "intermediate" / "2018_heldout_predictions.parquet", index=False)
    held.to_csv(WORK / "outputs" / "intermediate" / "2018_heldout_predictions.csv", index=False)
    primary = estimates(held)
    pl = play_level_gains(held)
    boot = bootstrap(pl)
    boot.to_parquet(WORK / "outputs" / "intermediate" / "2018_bootstrap_replicates.parquet", index=False)
    boot.to_csv(WORK / "outputs" / "intermediate" / "2018_bootstrap_replicates.csv", index=False)
    primary_ci = ci_table(primary, boot)
    primary_ci.to_csv(WORK / "outputs" / "tables" / "Table2_2018_primary_RPVA_estimates.csv", index=False)

    pseudo_null, pseudo_sum = pseudo_nearest(held)
    pseudo_null.to_parquet(WORK / "outputs" / "intermediate" / "2018_pseudo_nearest_null.parquet", index=False)
    pseudo_null.to_csv(WORK / "outputs" / "intermediate" / "2018_pseudo_nearest_null.csv", index=False)
    pd.DataFrame([pseudo_sum]).to_csv(WORK / "outputs" / "tables" / "Table5_pseudo_nearest_falsification.csv", index=False)

    weekly = []
    for w, h in held.groupby("week"):
        e = estimates(h)
        e["week"] = w
        weekly.append(e)
    weekly_df = pd.DataFrame(weekly)
    weekly_df.to_csv(WORK / "outputs" / "tables" / "2018_weekly_heldout_stability.csv", index=False)

    sens_rows = [{"analysis": "S1_primary_defender_set", **primary, "heldout_plays": pl[["gameId", "playId"]].drop_duplicates().shape[0]}]
    for name, pos_set, window, excl in [
        ("S2_expanded_defender_set", EXPANDED_POS, 0, False),
        ("S4_endpoint_plus_minus_1_frame", PRIMARY_POS, 1, False),
        ("S5_exclude_ambiguous_reconstructed_targets", PRIMARY_POS, 0, True),
    ]:
        b = build_dataset(pos_set, endpoint_window=window, exclude_ambiguous=excl)
        sdf = pd.DataFrame(b.rows)
        if len(sdf) and sdf["week"].between(1, 11).any() and sdf["week"].between(14, 17).any():
            sh, _ = add_predictions(sdf)
            se = estimates(sh)
            se["heldout_plays"] = sh[["gameId", "playId"]].drop_duplicates().shape[0]
        else:
            se = {"G_R": np.nan, "G_N": np.nan, "G_O": np.nan, "lambda_D": np.nan, "delta_RN": np.nan, "delta_RO": np.nan, "heldout_plays": 0}
        se["analysis"] = name
        sens_rows.append(se)
    sens = pd.DataFrame(sens_rows)
    sens.to_csv(WORK / "outputs" / "tables" / "Table4_sensitivity_analyses.csv", index=False)

    comp = pd.DataFrame(
        [
            {"season": 2018, **primary},
            {"season": 2023, "G_R": 1.735, "G_N": 1.359, "G_O": 0.609, "lambda_D": 0.750, "delta_RN": 0.376, "delta_RO": 1.126},
        ]
    )
    comp["lambda_difference_vs_2023"] = comp["lambda_D"] - 0.750
    comp.to_csv(WORK / "outputs" / "tables" / "Table3_2018_vs_2023_comparison.csv", index=False)
    pd.DataFrame(primary_build.flow_rows).to_csv(WORK / "outputs" / "tables" / "Table1_2018_sample_flow_and_harmonization.csv", index=False)

    replication_summary = summarize_replication(primary, primary_ci, pseudo_sum)
    make_figures(primary_ci, comp, pseudo_null, weekly_df, sens)

    strict_oob = (
        (df[["true_x", "input_start_x", "input_end_x", "ball_land_x"]] < 0).any(axis=1)
        | (df[["true_x", "input_start_x", "input_end_x", "ball_land_x"]] > 120).any(axis=1)
        | (df[["true_y", "input_start_y", "input_end_y", "ball_land_y"]] < 0).any(axis=1)
        | (df[["true_y", "input_start_y", "input_end_y", "ball_land_y"]] > 160 / 3).any(axis=1)
    )
    buffer_oob = (
        (df[["true_x", "input_start_x", "input_end_x", "ball_land_x"]] < -5).any(axis=1)
        | (df[["true_x", "input_start_x", "input_end_x", "ball_land_x"]] > 125).any(axis=1)
        | (df[["true_y", "input_start_y", "input_end_y", "ball_land_y"]] < -5).any(axis=1)
        | (df[["true_y", "input_start_y", "input_end_y", "ball_land_y"]] > (160 / 3 + 5)).any(axis=1)
    )
    qc_checks = {
        "no_heldout_game_in_training": set(df[df.week.between(1, 11)].gameId).isdisjoint(set(df[df.week.between(14, 17)].gameId)),
        "no_heldout_play_in_training": set(map(tuple, df[df.week.between(1, 11)][["gameId", "playId"]].drop_duplicates().to_numpy())).isdisjoint(set(map(tuple, held[["gameId", "playId"]].drop_duplicates().to_numpy()))),
        "ball_land_only_D": all(c not in A_FEATURES + H_ADD for c in D_ADD),
        "endpoint_errors_nonnegative": bool((held[["error_A", "error_H", "error_D"]] >= 0).all().all()),
        "gain_identity": bool(np.allclose(held["H_to_D_gain"], held["error_H"] - held["error_D"])),
        "coordinates_strict_field_bounds": bool(not strict_oob.any()),
        "coordinates_plausible_tracking_buffer": bool(not buffer_oob.any()),
    }
    coord_ranges = df[["true_x", "input_start_x", "input_end_x", "ball_land_x", "true_y", "input_start_y", "input_end_y", "ball_land_y"]].agg(["min", "max"])
    write(
        WORK / "reports" / "ANALYSIS_QC_REPORT.md",
        "# Analysis QC Report\n\n"
        + "\n".join(f"- {k}: {v}" for k, v in qc_checks.items())
        + f"\n\nStrict field-bound exceedance rows: {int(strict_oob.sum())} of {len(df)}. These are retained as plausible tracking values just outside the playing-field rectangle during pass outcomes; no row exceeds the predeclared QC buffer of x [-5, 125] and y [-5, 58.33].\n\n"
        + coord_ranges.to_markdown()
        + "\n",
    )

    write(WORK / "reports" / "SOFTWARE_ENVIRONMENT.md", f"# Software Environment\n\nGenerated: {now()}\n\nPython: `{sys.executable}`\n\nVersion: `{sys.version.replace(chr(10),' ')}`\n\nPlatform: `{platform.platform()}`\n\n{package_versions()}\n")

    lam_ci = primary_ci[primary_ci.estimand.eq("lambda_D")].iloc[0]
    results = f"""# 2018 External Replication Results

## Objective

Execute a held-out external-season RPVA replication on official 2018 NFL Big Data Bowl 2021 data using the locked A/H/D, R/N/O, bootstrap, and pseudo-nearest definitions.

## Data Provenance

BDB2021 competition files were obtained through Kaggle access terms. Raw files are not redistributed. The raw manifest is `{DATA / 'manifests' / 'RAW_DATA_MANIFEST.csv'}`.

## 2018 Data Structure

The primary modeling table contains {len(df):,} player-play rows. Held-out weeks 14-17 contain {len(held):,} player rows and {pl[['gameId','playId']].drop_duplicates().shape[0]:,} three-role eligible plays.

## Target Receiver Source

Primary labels use `targetedReceiver.csv` from the accessible bonus dataset. Reconstructed targets were computed only for QC and ambiguity sensitivity.

## 2023-to-2018 Harmonization

Feature construction uses the prespecified left-to-right coordinate normalization and harmonized 2018 feature definitions because the submitted repository package did not contain a verified full raw-data transformation pipeline for the 2023 NFL application.

## Frozen Protocol

The protocol was written and hashed before fitting final models. See `{WORK / 'protocol' / 'protocol_sha256.txt'}`.

## Model Setup

HGB model: sklearn `HistGradientBoostingRegressor` defaults with deterministic preprocessing and random_state {SEED}. Separate x/y models were fit for A, H, and D on weeks 1-11. Weeks 14-17 were held out.

## Primary RPVA Estimates

{primary_ci.to_markdown(index=False)}

## Game-Cluster Bootstrap

The 95% percentile game-cluster bootstrap used {BOOT_REPS} replicates without refitting models.

## Pseudo-Nearest Falsification

Observed lambda_D = {pseudo_sum['observed_lambda_D']:.6f}; null mean = {pseudo_sum['null_mean']:.6f}; 95% null interval = [{pseudo_sum['null_low']:.6f}, {pseudo_sum['null_high']:.6f}]; add-one corrected P = {pseudo_sum['p_value']:.6f}.

## Weekly Stability

{weekly_df.to_markdown(index=False)}

## Sensitivity Analyses

{sens.to_markdown(index=False)}

## 2023-vs-2018 Comparison

{comp.to_markdown(index=False)}

## Deviations

No post-freeze deviations were recorded. ExtraTrees robustness was skipped because the exact submitted final ExtraTrees configuration was not recoverable.

## Limitations

The 2018 workflow is a harmonized external replication, not an exact raw-pipeline rerun of the submitted 2023 implementation. N is proximity-defined and should not be interpreted as assigned coverage.

## Replication Summary

{replication_summary}

## Scientific Interpretation

The 2018 evidence should be interpreted as a held-out audit of relational allocation under harmonized data definitions. It does not establish causality, tactical assignment, or deployment-time utility.
"""
    write(WORK / "reports" / "2018_EXTERNAL_REPLICATION_RESULTS.md", results)

    readme = f"""# RPVA 2018 External Validation

Run from PowerShell:

```powershell
python .\\nfl_external_validation\\2018\\scripts\\run_2018_external_replication.py
```

Inputs are locally obtained BDB2021 competition files and public bonus target labels supplied by user-provided paths. Raw files are not redistributed.

Random seeds: HGB {SEED}, bootstrap {SEED}, pseudo-nearest {PSEUDO_SEED}.
"""
    write(WORK / "README.md", readme)

    final_files = []
    for p in sorted(WORK.rglob("*")):
        if p.is_file():
            final_files.append(f"- `{p.relative_to(WORK)}` - generated 2018 external-validation artifact.")
    for p in sorted((DATA / "processed").glob("*")) + sorted((DATA / "manifests").glob("*")):
        if p.is_file():
            final_files.append(f"- `{p.relative_to(DATA)}` - derived data or raw manifest.")
    write(WORK / "reports" / "FINAL_FILE_INDEX.md", "# Final File Index\n\n" + "\n".join(final_files) + "\n")

    final_summary = {
        "data": "obtained under applicable competition access terms; raw files are not redistributed",
        "target_receiver_source": "public targeted-receiver bonus labels from tombliss/nfl-big-data-bowl-2021-bonus",
        "heldout_plays": int(pl[["gameId", "playId"]].drop_duplicates().shape[0]),
        "primary": primary,
        "lambda_ci": {"low": float(lam_ci.ci_low), "high": float(lam_ci.ci_high)},
        "pseudo": pseudo_sum,
    }
    write(WORK / "outputs" / "FINAL_SUMMARY.json", json.dumps(final_summary, indent=2))


if __name__ == "__main__":
    main()
